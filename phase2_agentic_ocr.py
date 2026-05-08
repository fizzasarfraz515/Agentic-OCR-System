# # ============================================================
# # PHASE 2 — PRIVACY-AWARE AGENTIC OCR SYSTEM
# # Professional Practices in IT — Final Project
# #
# # FEATURES:
# # ✅ Agentic OCR Pipeline
# # ✅ Confidential Data Detection
# # ✅ Auto Redaction
# # ✅ Human-in-the-Loop
# # ✅ Ethical AI
# # ✅ GDPR + PECA 2016 Awareness
# # ✅ Secure Memory
# # ✅ Word Export
# # ✅ Multi-language OCR
# # ============================================================

# import streamlit as st
# from PIL import Image, ImageEnhance, ImageFilter, ImageOps
# import pytesseract
# from docx import Document
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# import io
# import platform
# import time
# import hashlib
# import datetime
# import re
# import statistics

# # ============================================================
# # TESSERACT SETUP
# # ============================================================

# if platform.system() == "Windows":
#     pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # ============================================================
# # PAGE CONFIG
# # ============================================================

# st.set_page_config(
#     page_title="Privacy-Aware Agentic OCR",
#     page_icon="🔒",
#     layout="wide"
# )

# # ============================================================
# # CUSTOM CSS
# # ============================================================

# st.markdown("""
# <style>

# .main {
#     background: linear-gradient(to right, #0f172a, #1e293b);
# }

# .block-container {
#     padding-top: 2rem;
# }

# .agent-box {
#     background: rgba(255,255,255,0.05);
#     padding: 20px;
#     border-radius: 15px;
#     border: 1px solid rgba(255,255,255,0.1);
#     margin-bottom: 20px;
# }

# .log-box {
#     background: black;
#     color: #00ff99;
#     padding: 10px;
#     border-radius: 10px;
#     font-family: monospace;
#     font-size: 0.8rem;
# }

# .safe-box {
#     background: rgba(0,255,100,0.08);
#     border: 1px solid rgba(0,255,100,0.3);
#     padding: 15px;
#     border-radius: 10px;
# }

# .warning-box {
#     background: rgba(255,0,0,0.08);
#     border: 1px solid rgba(255,0,0,0.3);
#     padding: 15px;
#     border-radius: 10px;
# }

# </style>
# """, unsafe_allow_html=True)

# # ============================================================
# # SESSION MEMORY
# # ============================================================

# if "logs" not in st.session_state:
#     st.session_state.logs = []

# if "history" not in st.session_state:
#     st.session_state.history = []

# if "stats" not in st.session_state:
#     st.session_state.stats = {
#         "processed": 0,
#         "avg_confidence": 0,
#         "words": 0
#     }

# # ============================================================
# # LOGGING
# # ============================================================

# def log(message):
#     timestamp = datetime.datetime.now().strftime("%H:%M:%S")
#     st.session_state.logs.append(f"[{timestamp}] {message}")

# # ============================================================
# # CONFIDENTIAL DATA PATTERNS
# # ============================================================

# SENSITIVE_PATTERNS = {

#     # Emails
#     "Email":
#         r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",

#     # Phone numbers
#     "Phone Number":
#         r"\+?\d[\d\s\-]{7,15}",

#     # Credit Cards
#     "Credit Card":
#         r"\b(?:\d[ -]*?){13,16}\b",

#     # IBAN
#     "IBAN":
#         r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b",

#     # Passport
#     "Passport":
#         r"\b[A-Z0-9]{6,9}\b",

#     # SSN (US)
#     "SSN":
#         r"\b\d{3}-\d{2}-\d{4}\b",

#     # CNIC
#     "CNIC":
#         r"\b\d{5}-\d{7}-\d\b",

#     # Driver License
#     "Driver License":
#         r"\b[A-Z]{1,2}\d{6,12}\b",

#     # Passwords
#     "Password":
#         r"(?i)\b(password|passwd|pin|secret|api key|token)\b",

#     # Bank Account
#     "Bank Account":
#         r"\b\d{9,18}\b",

#     # Crypto Wallet
#     "Crypto Wallet":
#         r"\b[13][a-km-zA-HJ-NP-Z1-9]{25,34}\b",

#     # Aadhaar India
#     "Aadhaar":
#         r"\b\d{4}\s\d{4}\s\d{4}\b",

#     # IP Address
#     "IP Address":
#         r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b",
# }

# # ============================================================
# # IMAGE ANALYSIS AGENT
# # ============================================================

# def analyze_image(image):

#     gray = image.convert("L")

#     brightness = statistics.mean(gray.getdata())

#     decisions = {
#         "grayscale": True,
#         "contrast": True,
#         "sharpen": True,
#         "threshold": True,
#         "brightness": brightness
#     }

#     log(f"Image brightness analyzed: {brightness:.1f}")

#     return decisions

# # ============================================================
# # PREPROCESSING AGENT
# # ============================================================

# def preprocess_image(image, decisions):

#     steps = []

#     img = image.copy()

#     if decisions["grayscale"]:
#         img = img.convert("L")
#         steps.append("Grayscale conversion")

#     if decisions["contrast"]:
#         img = ImageEnhance.Contrast(img).enhance(2)
#         steps.append("Contrast enhancement")

#     if decisions["sharpen"]:
#         img = img.filter(ImageFilter.SHARPEN)
#         steps.append("Sharpening")

#     if decisions["threshold"]:
#         img = img.point(lambda x: 0 if x < 128 else 255)
#         steps.append("Threshold binarization")

#     log("Preprocessing complete")

#     return img, steps

# # ============================================================
# # OCR AGENT
# # ============================================================

# def extract_text(img, lang="eng"):

#     try:

#         data = pytesseract.image_to_data(
#             img,
#             output_type=pytesseract.Output.DICT
#         )

#         confs = []

#         for c in data["conf"]:
#             try:
#                 c = int(c)
#                 if c >= 0:
#                     confs.append(c)
#             except:
#                 pass

#         confidence = statistics.mean(confs) if confs else 0

#         text = pytesseract.image_to_string(img, lang=lang)

#         log(f"OCR extraction complete: confidence={confidence:.1f}")

#         return text, confidence

#     except Exception as e:
#         log(f"OCR error: {e}")
#         return "", 0

# # ============================================================
# # TEXT CLEANING AGENT
# # ============================================================

# def clean_text(text):

#     text = re.sub(r"\n{3,}", "\n\n", text)

#     text = re.sub(r"[ ]{2,}", " ", text)

#     text = text.strip()

#     log("Post-processing complete")

#     return text

# # ============================================================
# # CONFIDENTIALITY AGENT
# # ============================================================

# def confidentiality_agent(text):

#     findings = []

#     redacted_text = text

#     sensitivity_score = 0

#     for label, pattern in SENSITIVE_PATTERNS.items():

#         matches = re.findall(pattern, text)

#         if matches:

#             findings.append({
#                 "type": label,
#                 "count": len(matches)
#             })

#             sensitivity_score += len(matches)

#             redacted_text = re.sub(
#                 pattern,
#                 f"[REDACTED_{label}]",
#                 redacted_text
#             )

#     log(f"Confidential scan complete: {sensitivity_score} sensitive items")

#     return redacted_text, findings, sensitivity_score

# # ============================================================
# # ETHICAL AI AGENT
# # ============================================================

# def ethical_check(confidence, findings):

#     warnings = []

#     if confidence < 50:
#         warnings.append(
#             "Low OCR confidence detected. Verify extracted text manually."
#         )

#     if findings:
#         warnings.append(
#             "Sensitive information detected. Handle responsibly."
#         )

#     return warnings

# # ============================================================
# # SECURE MEMORY AGENT
# # ============================================================

# def update_memory(confidence, word_count, file_hash):

#     stats = st.session_state.stats

#     stats["processed"] += 1
#     stats["words"] += word_count

#     n = stats["processed"]

#     stats["avg_confidence"] = (
#         (
#             stats["avg_confidence"] * (n - 1)
#         ) + confidence
#     ) / n

#     secure_hash = hashlib.sha256(
#         file_hash.encode()
#     ).hexdigest()

#     st.session_state.history.append({
#         "hash": secure_hash,
#         "confidence": confidence,
#         "words": word_count,
#         "time": datetime.datetime.now().isoformat()
#     })

#     log("Secure memory updated")

# # ============================================================
# # WORD EXPORT AGENT
# # ============================================================

# def generate_doc(text, confidence, findings):

#     doc = Document()

#     title = doc.add_heading(
#         "Privacy-Aware Agentic OCR Export",
#         0
#     )

#     title.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     doc.add_heading("Metadata", level=1)

#     doc.add_paragraph(f"Confidence: {confidence:.1f}%")
#     doc.add_paragraph(f"Generated: {datetime.datetime.now()}")

#     doc.add_heading("Privacy Report", level=1)

#     if findings:
#         for f in findings:
#             doc.add_paragraph(
#                 f"{f['type']} detected ({f['count']})"
#             )
#     else:
#         doc.add_paragraph("No sensitive information detected.")

#     doc.add_heading("Extracted Text", level=1)

#     doc.add_paragraph(text)

#     doc.add_heading("Legal Notice", level=1)

#     p = doc.add_paragraph(
#         "This document was processed locally. "
#         "Users are responsible for handling "
#         "personal data according to GDPR, PECA 2016, "
#         "and applicable privacy laws."
#     )

#     p.runs[0].font.size = Pt(9)
#     p.runs[0].font.color.rgb = RGBColor(120,120,120)

#     buffer = io.BytesIO()

#     doc.save(buffer)

#     buffer.seek(0)

#     return buffer

# # ============================================================
# # SIDEBAR
# # ============================================================

# with st.sidebar:

#     st.title("🤖 Agent Control")

#     consent = st.checkbox(
#         "I consent to local processing"
#     )

#     privacy = st.checkbox(
#         "I understand privacy responsibilities"
#     )

#     language = st.selectbox(
#         "OCR Language",
#         [
#             "eng",
#             "urd",
#             "ara",
#             "fra",
#             "deu"
#         ]
#     )

#     autonomy = st.radio(
#         "Autonomy Mode",
#         [
#             "Semi-Autonomous",
#             "Full Autonomous"
#         ]
#     )

#     st.markdown("---")

#     st.metric(
#         "Processed",
#         st.session_state.stats["processed"]
#     )

#     st.metric(
#         "Avg Confidence",
#         f"{st.session_state.stats['avg_confidence']:.1f}%"
#     )

#     st.metric(
#         "Words",
#         st.session_state.stats["words"]
#     )

# # ============================================================
# # MAIN UI
# # ============================================================

# st.title("🔒 Privacy-Aware Agentic OCR System")

# st.write("""
# This OCR system acts as an intelligent privacy-aware AI agent capable of:

# ✅ OCR Extraction  
# ✅ Confidential Data Detection  
# ✅ Automatic Redaction  
# ✅ Ethical AI Warnings  
# ✅ Human-in-the-Loop Approval  
# ✅ Secure Memory  
# ✅ Privacy Protection  
# """)

# # ============================================================
# # FILE UPLOAD
# # ============================================================

# if not consent or not privacy:

#     st.warning(
#         "Please accept consent and privacy requirements."
#     )

#     st.stop()

# uploaded_file = st.file_uploader(
#     "Upload Image",
#     type=["png", "jpg", "jpeg", "bmp", "tiff"]
# )

# # ============================================================
# # MAIN PIPELINE
# # ============================================================

# if uploaded_file:

#     image = Image.open(uploaded_file)

#     st.image(image, caption="Uploaded Image")

#     if st.button("🚀 Run Agentic OCR"):

#         # ====================================================
#         # PERCEPTION
#         # ====================================================

#         log("PERCEPTION started")

#         file_bytes = uploaded_file.getvalue()

#         file_hash = hashlib.md5(file_bytes).hexdigest()

#         # ====================================================
#         # DECISION
#         # ====================================================

#         decisions = analyze_image(image)

#         # ====================================================
#         # ACTION
#         # ====================================================

#         processed_img, steps = preprocess_image(
#             image,
#             decisions
#         )

#         raw_text, confidence = extract_text(
#             processed_img,
#             language
#         )

#         cleaned_text = clean_text(raw_text)

#         # ====================================================
#         # CONFIDENTIALITY AGENT
#         # ====================================================

#         safe_text, findings, sensitivity_score = confidentiality_agent(
#             cleaned_text
#         )

#         # ====================================================
#         # ETHICAL AGENT
#         # ====================================================

#         warnings = ethical_check(
#             confidence,
#             findings
#         )

#         # ====================================================
#         # LEARNING
#         # ====================================================

#         update_memory(
#             confidence,
#             len(cleaned_text.split()),
#             file_hash
#         )

#         # ====================================================
#         # RESULTS
#         # ====================================================

#         st.success("Extraction Complete")

#         col1, col2 = st.columns([2,1])

#         with col1:

#             st.subheader("📝 Extracted Text")

#             if findings:

#                 st.markdown("""
#                 <div class="warning-box">
#                 Sensitive information detected.
#                 Text has been automatically redacted.
#                 </div>
#                 """, unsafe_allow_html=True)

#                 reveal = st.checkbox(
#                     "Show original unredacted text"
#                 )

#                 if reveal:
#                     display_text = cleaned_text
#                 else:
#                     display_text = safe_text

#             else:
#                 display_text = cleaned_text

#             st.text_area(
#                 "",
#                 display_text,
#                 height=350
#             )

#         with col2:

#             st.subheader("📊 Results")

#             st.metric(
#                 "Confidence",
#                 f"{confidence:.1f}%"
#             )

#             st.metric(
#                 "Words",
#                 len(cleaned_text.split())
#             )

#             st.metric(
#                 "Privacy Risk",
#                 sensitivity_score
#             )

#             st.metric(
#                 "Sensitive Types",
#                 len(findings)
#             )

#         # ====================================================
#         # CONFIDENTIAL FINDINGS
#         # ====================================================

#         if findings:

#             st.subheader("🔒 Confidential Information Report")

#             for item in findings:

#                 st.warning(
#                     f"{item['type']} detected "
#                     f"({item['count']} occurrence(s))"
#                 )

#         # ====================================================
#         # ETHICAL WARNINGS
#         # ====================================================

#         if warnings:

#             st.subheader("⚖️ Ethical Warnings")

#             for w in warnings:

#                 st.warning(w)

#         # ====================================================
#         # PREPROCESSING REPORT
#         # ====================================================

#         st.subheader("⚙️ Agent Decisions")

#         for s in steps:
#             st.write(f"• {s}")

#         # ====================================================
#         # LOGS
#         # ====================================================

#         st.subheader("🧠 Agent Logs")

#         for l in reversed(st.session_state.logs[-20:]):

#             st.markdown(
#                 f"""
#                 <div class="log-box">{l}</div>
#                 """,
#                 unsafe_allow_html=True
#             )

#         # ====================================================
#         # EXPORT
#         # ====================================================

#         doc = generate_doc(
#             display_text,
#             confidence,
#             findings
#         )

#         st.download_button(
#             "⬇ Download Safe Word Document",
#             data=doc,
#             file_name="privacy_safe_ocr.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

# # ============================================================
# # FOOTER
# # ============================================================

# st.markdown("---")

# st.markdown("""
# ### 📋 System Architecture

# Perception → Decision → OCR → Confidentiality Agent →
# Ethical AI → Human Approval → Secure Export → Learning

# ### ⚖️ Compliance

# ✅ GDPR Mindset  
# ✅ PECA 2016  
# ✅ ACM Code of Ethics  
# ✅ IEEE Ethical AI Principles  
# ✅ Human-in-the-Loop AI  
# """)


# # ================================
# # 🔐 Privacy OCR Shield (PRO VERSION)
# # ================================

# import streamlit as st
# from PIL import Image, ImageEnhance, ImageFilter
# import pytesseract
# import re
# import datetime
# import io
# from docx import Document

# # ────────────────────────────────
# # TESSERACT PATH (Windows)
# # ────────────────────────────────
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# # ────────────────────────────────
# # PAGE CONFIG
# # ────────────────────────────────
# st.set_page_config(
#     page_title="🔐 Agentic OCR system",
#     page_icon="🔐",
#     layout="wide"
# )

# # ────────────────────────────────
# # UI THEME (BEIGE + CLEAN DASHBOARD)
# # ────────────────────────────────
# st.markdown("""
# <style>
# .stApp { background-color: #f3eadc !important; }

# .card {
#     background: #fff7ec;
#     padding: 18px;
#     border-radius: 14px;
#     border: 1px solid #e7d6bd;
#     box-shadow: 0 6px 18px rgba(0,0,0,0.05);
#     margin-bottom: 10px;
# }

# .badge {
#     padding: 6px 12px;
#     border-radius: 10px;
#     background: #d9c2a7;
#     font-weight: bold;
# }

# </style>
# """, unsafe_allow_html=True)

# st.title("🔐 Agentic OCR system")

# # ────────────────────────────────
# # CONFIDENTIAL PATTERNS (EXPANDED)
# # ────────────────────────────────
# PATTERNS = {
#     "CNIC": (r"\b\d{5}-\d{7}-\d{1}\b", "█████-███████-█"),
#     "PASSPORT": (r"\b[A-Z]{1,2}[0-9]{6,9}\b", "██████████"),
#     "EMAIL": (r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", "████@████.██"),
#     "PHONE": (r"\b03\d{2}-?\d{7}\b", "03██-███████"),
#     "CARD": (r"\b(?:\d{4}[- ]?){3}\d{4}\b", "████-████-████-████"),
#     "IBAN": (r"\b[A-Z]{2}\d{2}[A-Z0-9]{10,30}\b", "████████████"),
# }

# # ────────────────────────────────
# # IMAGE PREPROCESSING (FIX BLUR OCR)
# # ────────────────────────────────
# def enhance_image(img):
#     img = img.convert("L")  # grayscale
#     img = ImageEnhance.Contrast(img).enhance(2)
#     img = ImageEnhance.Sharpness(img).enhance(2)
#     return img


# # ────────────────────────────────
# # OCR ENGINE (SAFE + CONFIDENCE)
# # ────────────────────────────────
# def extract_text(img):
#     try:
#         data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

#         words = []
#         confidences = []

#         for i in range(len(data["text"])):
#             word = data["text"][i].strip()
#             conf = int(data["conf"][i]) if data["conf"][i] != "-1" else -1

#             if word and len(word) > 1:
#                 words.append(word)
#                 if conf > 0:
#                     confidences.append(conf)

#         text = " ".join(words)

#         avg_conf = sum(confidences) / len(confidences) if confidences else 0

#         return text, avg_conf

#     except Exception as e:
#         return "", 0


# # ────────────────────────────────
# # CLEAN GARBAGE OCR
# # ────────────────────────────────
# def clean_text(text):
#     text = re.sub(r"[^\x00-\x7F]+", " ", text)

#     words = text.split()
#     cleaned = []

#     for w in words:
#         if len(w) <= 2 and not w.isdigit():
#             continue
#         if re.fullmatch(r"[^\w]+", w):
#             continue
#         cleaned.append(w)

#     return " ".join(cleaned)


# # ────────────────────────────────
# # SHIELD ENGINE
# # ────────────────────────────────
# def shield(text):
#     findings = []
#     redacted = text
#     risk = 0

#     for label, (pattern, mask) in PATTERNS.items():
#         matches = re.findall(pattern, text)

#         risk += len(matches) * 10

#         for m in matches:
#             findings.append((label, m))

#         redacted = re.sub(pattern, mask, redacted)

#     return redacted, findings, risk


# # ────────────────────────────────
# # QUALITY SCORE
# # ────────────────────────────────
# def quality_score(conf, length):
#     if length == 0:
#         return 0
#     return min(100, conf)


# # ────────────────────────────────
# # FILE UPLOAD
# # ────────────────────────────────
# file = st.file_uploader("📤 Upload Document", type=["png", "jpg", "jpeg"])

# if file:

#     img = Image.open(file)
#     img = enhance_image(img)

#     col1, col2 = st.columns(2)

#     with col1:
#         st.markdown('<div class="card">', unsafe_allow_html=True)
#         st.subheader("📄 Document Preview")
#         st.image(img, use_container_width=True)
#         st.markdown('</div>', unsafe_allow_html=True)

#     # OCR
#     raw_text, conf = extract_text(img)
#     clean = clean_text(raw_text)

#     redacted, findings, risk = shield(clean)

#     qscore = quality_score(conf, len(clean))

#     # ───── DASHBOARD
#     with col2:
#         st.markdown('<div class="card">', unsafe_allow_html=True)
#         st.subheader("📊 AI Dashboard")

#         st.metric("OCR Confidence", f"{conf:.2f}%")
#         st.metric("Quality Score", f"{qscore:.2f}/100")
#         st.metric("Risk Score", risk)
#         st.metric("Sensitive Information", len(findings))

#         status = "SAFE" if risk == 0 else "REDACTED"
#         st.markdown(f"<div class='badge'>{status}</div>", unsafe_allow_html=True)

#         st.markdown('</div>', unsafe_allow_html=True)

#     # ───── TABS
#     tab1, tab2, tab3 = st.tabs(["🧾 OCR", "🔐 Redacted", "🚨 Findings"])

#     with tab1:
#         st.code(clean)

#     with tab2:
#         st.code(redacted)

#     with tab3:
#         if findings:
#             for f in findings:
#                 st.error(f"{f[0]} → {f[1]}")
#         else:
#             st.success("No sensitive data detected")


#     # ───── DOWNLOADS
#     st.markdown("---")

#     report = f"""
# OCR REPORT
# Time: {datetime.datetime.now()}

# Confidence: {conf:.2f}
# Quality: {qscore:.2f}
# Risk: {risk}

# RAW:
# {clean}

# REDACTED:
# {redacted}

# FINDINGS:
# {findings}
# """

#     col1, col2 = st.columns(2)

#     with col1:
#         st.download_button("⬇️ TXT Report", report, "report.txt")

#     with col2:
#         doc = Document()
#         doc.add_paragraph(report)

#         buffer = io.BytesIO()
#         doc.save(buffer)
#         buffer.seek(0)

#         st.download_button(
#             "⬇️ DOCX Report",
#             buffer,
#             "report.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

# # ────────────────────────────────
# # CONSENT (BOTTOM ONLY)
# # ────────────────────────────────
# st.markdown("---")
# consent = st.checkbox("✔ I confirm I have permission to process this document")

# if not consent:
#     st.warning("Please accept consent to enable processing")





import streamlit as st
from PIL import Image
import pytesseract
import re
import datetime
import io
from docx import Document

# ────────────────────────────────
# TESSERACT CONFIG
# ────────────────────────────────
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ────────────────────────────────
# PAGE CONFIG
# ────────────────────────────────
st.set_page_config(
    page_title="🔐 AGENTIC OCR SYSTEM",
    page_icon="🔐",
    layout="wide"
)

# ────────────────────────────────
# BEIGE THEME (FULL CLEAN DESIGN)
# ────────────────────────────────
st.markdown("""
<style>

/* GLOBAL BACKGROUND */
body, .main {
    background-color: #f3eadb !important;
    color: #3b2f24;
}

/* TITLE */
.title {
    font-size: 40px;
    font-weight: 800;
    color: #4a3a2a;
    margin-bottom: 10px;
}

/* SIDEBAR */
[data-testid="stSidebar"] {
    background-color: #eadcc9;
}
[data-testid="stSidebar"] * {
    color: #3b2f24 !important;
}

/* CARDS */
.card {
    background-color: #fff7ec;
    padding: 20px;
    border-radius: 16px;
    border: 1px solid #e2c9a6;
    box-shadow: 0px 3px 10px rgba(0,0,0,0.06);
}

/* METRICS */
.metric-box {
    background-color: #f7ead7;
    padding: 15px;
    border-radius: 12px;
    text-align: center;
    border: 1px solid #e2c9a6;
}

/* CODE BLOCK */
pre {
    background-color: #fff3e3 !important;
    border-radius: 10px;
    padding: 10px;
}

/* FOOTER */
.footer {
    text-align: center;
    color: #6b5845;
    margin-top: 40px;
    font-size: 13px;
}

</style>
""", unsafe_allow_html=True)

# ────────────────────────────────
# HEADER
# ────────────────────────────────
st.markdown('<div class="title">🔐 AGENTIC OCR SYSTEM</div>', unsafe_allow_html=True)
st.caption("Beige Secure Document Analyzer — CNIC • Passport • Email Protection")

# ────────────────────────────────
# PRIVACY PATTERNS
# ────────────────────────────────
PATTERNS = {
    "CNIC": (r"\b\d{5}-\d{7}-\d{1}\b", "█████-███████-█"),
    "PASSPORT": (r"\b[A-Z]{1,2}\d{6,8}\b", "██████████"),
    "EMAIL": (r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", "████@████.██"),
    "PHONE": (r"\b03\d{2}-?\d{7}\b", "03██-███████"),
    "CARD": (r"\b(?:\d{4}[- ]?){3}\d{4}\b", "████-████-████-████"),
}

# ────────────────────────────────
# SIDEBAR
# ────────────────────────────────
st.sidebar.markdown("## ⚙️ Privacy Controls")

consent = st.sidebar.checkbox("✔ I have consent to process this document")
mode = st.sidebar.selectbox("Mode", ["Secure Redaction", "Detection Only"])

if not consent:
    st.warning("⚠️ Consent required to continue")
    st.stop()

# ────────────────────────────────
# OCR FUNCTION
# ────────────────────────────────
def extract_text(img):
    text = pytesseract.image_to_string(img, config="--psm 3")
    text = re.sub(r"\s+", " ", text)
    return text.strip()

# ────────────────────────────────
# PRIVACY SHIELD ENGINE
# ────────────────────────────────
def shield(text):
    findings = []
    redacted = text

    for label, (pattern, mask) in PATTERNS.items():
        matches = re.findall(pattern, redacted)

        for m in matches:
            findings.append((label, m))

        redacted = re.sub(pattern, mask, redacted)

    return redacted, findings

# ────────────────────────────────
# DOC EXPORT (SAFE ONLY)
# ────────────────────────────────
def make_doc(redacted, findings):
    doc = Document()
    doc.add_heading("Privacy OCR Report", 0)
    doc.add_paragraph(str(datetime.datetime.now()))

    doc.add_heading("Redacted Output (SAFE)", level=1)
    doc.add_paragraph(redacted)

    doc.add_heading("Findings", level=1)

    if findings:
        for f in findings:
            doc.add_paragraph(f"{f[0]} detected (hidden for privacy)")
    else:
        doc.add_paragraph("No sensitive data found")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ────────────────────────────────
# UPLOAD
# ────────────────────────────────
file = st.file_uploader("📤 Upload Document", type=["png", "jpg", "jpeg"])

if file:

    img = Image.open(file)

    raw = extract_text(img)
    redacted, findings = shield(raw)

    col1, col2 = st.columns(2)

    # ───────── IMAGE ─────────
    with col1:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.subheader("📄 Document Preview")
        st.image(img, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

    # ───────── DASHBOARD ─────────
    with col2:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.subheader("📊 Security Dashboard")

        colA, colB, colC = st.columns(3)

        with colA:
            st.markdown(f'<div class="metric-box"><h3>{len(findings)}</h3><p>Leaks Found</p></div>', unsafe_allow_html=True)

        with colB:
            status = "SAFE" if len(findings) == 0 else "REDACTED"
            st.markdown(f'<div class="metric-box"><h3>{status}</h3><p>Status</p></div>', unsafe_allow_html=True)

        with colC:
            st.markdown(f'<div class="metric-box"><h3>{mode}</h3><p>Mode</p></div>', unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)

    # ────────────────────────────────
    # OUTPUT TABS
    # ────────────────────────────────
    tab1, tab2, tab3 = st.tabs(["🔐 OCR Text", "🛡️ Redacted Output", "🚨 Findings"])

    with tab1:
        st.code(raw)

    with tab2:
        st.success("Secure Output")
        st.code(redacted)

    with tab3:
        if findings:
            for f in findings:
                st.error(f"{f[0]} → detected and removed")
        else:
            st.success("No sensitive data detected")

    # ────────────────────────────────
    # DOWNLOADS
    # ────────────────────────────────
    st.markdown("---")
    st.subheader("📥 Export Reports")

    txt = f"""
PRIVACY OCR REPORT
{datetime.datetime.now()}

RAW TEXT:
{raw}

REDACTED TEXT:
{redacted}

TOTAL LEAKS:
{len(findings)}
"""

    col1, col2 = st.columns(2)

    with col1:
        st.download_button("⬇️ TXT Report", txt, "privacy_report.txt")

    with col2:
        doc = make_doc(redacted, findings)
        st.download_button(
            "⬇️ DOCX Report",
            doc,
            "privacy_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ────────────────────────────────
# FOOTER
# ────────────────────────────────
st.markdown("""
<div class="footer">
🔐 Privacy OCR Shield • Beige Secure System • Agentic Document Intelligence
</div>
""", unsafe_allow_html=True)